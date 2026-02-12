import os
import sys
import cv2
import numpy as np
from paddleocr import PPStructure, PaddleOCR
from paddleocr.ppstructure.recovery.recovery_to_doc import convert_info_docx
from docx import Document
from docx.shared import Pt, Inches, Emu

class LocalOCREngine:
    def __init__(self, use_gpu=False, lang='ch'):
        self.use_gpu = use_gpu
        self.lang = lang
        
        # Initialize PaddleOCR (Fallback to non-structure engine to fix crash)
        print(f"Initializing PaddleOCR (v4) with lang={lang}...")
        try:
            self.table_engine = PaddleOCR(
                show_log=True, 
                use_gpu=use_gpu, 
                lang=lang,
                ocr_version='PP-OCRv4',
                use_angle_cls=False
            )
            print(f"✓ Initialized PaddleOCR (v4) successfully.")
        except Exception as e:
            print(f"PaddleOCR init failed: {e}")
            raise e
        
    def custom_convert_to_docx(self, processed_output, save_path, img_width_px, img_height_px):
        """
        Generates a Word document where each text block is positioned 
        to match the original image layout.
        """
        doc = Document()
        
        # A4 Page dimensions in EMU (roughly 8.27in x 11.69in)
        # 1 inch = 914400 EMU
        page_width_emu = int(8.27 * 914400)
        page_height_emu = int(11.69 * 914400)
        
        # Scale factor from pixels to EMU
        scale_x = page_width_emu / img_width_px
        scale_y = page_height_emu / img_height_px

        # Collect all lines with their global coordinates for sorting
        all_lines = []
        for item in processed_output:
            for line in item.get('lines', []):
                if line.get('bbox'):
                    all_lines.append({
                        'text': line['text'],
                        'bbox': line['bbox']
                    })

        # Sort all lines by Y then X to maintain natural reading order flow with spacing
        all_lines.sort(key=lambda x: (x['bbox'][1], x['bbox'][0]))
        
        last_y = 0
        for line in all_lines:
            bbox = line['bbox']
            # x1, y1, x2, y2
            x_emu = int(bbox[0] * scale_x)
            y_emu = int(bbox[1] * scale_y)
            
            # Calculate delta Y for "space_before"
            delta_y = max(0, y_emu - last_y)
            
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Emu(x_emu)
            # We use space_before to simulate the vertical position
            # Note: space_before is slightly capped in Word UI but works well for layout
            # Cap space_before to avoid huge gaps failing docx
            p.paragraph_format.space_before = Emu(min(delta_y, 1000000)) 
            
            run = p.add_run(line['text'])
            run.font.size = Pt(9)
            run.font.name = 'Arial'
            
            # Update last_y based on the bottom of the current line
            last_y = int(bbox[3] * scale_y)

        doc.save(save_path)
        return save_path

    def process_image(self, img_path_or_array, save_folder="./output", img_name="result"):
        if not os.path.exists(save_folder):
            os.makedirs(save_folder)

        if isinstance(img_path_or_array, str):
            img = cv2.imread(img_path_or_array)
            img_name = os.path.basename(img_path_or_array).split('.')[0]
        else:
            img = img_path_or_array

        if img is None:
            raise ValueError("Image could not be loaded.")

        h, w = img.shape[:2]

        # Run the engine
        result = self.table_engine(img)
        
        # Compatibility Fix: Convert PaddleOCR format to PPStructure format if needed
        print(f"DEBUG: result type = {type(result)}")
        
        # PaddleOCR returns different formats:
        # - When called directly: returns a list [[box, (text, conf)], ...]
        # - When using certain modes: returns tuple (result_list, ...)
        if isinstance(result, tuple) and len(result) > 0:
            print(f"DEBUG: PaddleOCR returned tuple with {len(result)} elements")
            print(f"DEBUG: result[0] type = {type(result[0])}, length = {len(result[0]) if hasattr(result[0], '__len__') else 'N/A'}")
            
            # If first element is a list, it's likely the OCR results
            if isinstance(result[0], list) and len(result[0]) > 0:
                # Check if it's already in the correct format [box, (text, conf)]
                first_item = result[0][0]
                print(f"DEBUG: first_item type = {type(first_item)}")
                
                if isinstance(first_item, (list, tuple)) and len(first_item) >= 2:
                    # Format: [[box, (text, conf)], ...]
                    lines = result[0]
                    print(f"DEBUG: Found {len(lines)} lines with box+text format")
                    
                    formatted_res = []
                    for i, line in enumerate(lines):
                        if len(line) >= 2:
                            box = line[0]
                            text_info = line[1]
                            text = text_info[0] if isinstance(text_info, (list, tuple)) else str(text_info)
                            
                            pts = np.array(box)
                            x1, y1 = np.min(pts, axis=0)
                            x2, y2 = np.max(pts, axis=0)
                            formatted_res.append({
                                'text': text,
                                'bbox': [float(x1), float(y1), float(x2), float(y2)]
                            })
                            if i < 3:  # Log first 3 for debugging
                                print(f"DEBUG: Line {i}: text='{text}', bbox=[{x1:.1f},{y1:.1f},{x2:.1f},{y2:.1f}]")
                    
                    print(f"DEBUG: Converted to {len(formatted_res)} formatted results")
                    result = [{
                        'type': 'text',
                        'bbox': [0, 0, w, h],
                        'res': formatted_res
                    }]
                elif isinstance(first_item, np.ndarray):
                    # Format: tuple([box_array, box_array, ...], [text_info, text_info, ...], [scores, ...])
                    # Boxes and texts are in SEPARATE lists - need to zip them!
                    print(f"DEBUG: Boxes and texts are separate - zipping them together")
                    
                    boxes = result[0]  # List of box arrays
                    texts_and_scores = result[1] if len(result) > 1 else []  # List of (text, score) tuples
                    
                    print(f"DEBUG: Found {len(boxes)} boxes and {len(texts_and_scores)} text items")
                    
                    formatted_res = []
                    for i, (box, text_info) in enumerate(zip(boxes, texts_and_scores)):
                        text = text_info[0] if isinstance(text_info, (list, tuple)) else str(text_info)
                        
                        pts = np.array(box)
                        x1, y1 = np.min(pts, axis=0)
                        x2, y2 = np.max(pts, axis=0)
                        formatted_res.append({
                            'text': text,
                            'bbox': [float(x1), float(y1), float(x2), float(y2)]
                        })
                        if i < 3:  # Log first 3
                            print(f"DEBUG: Line {i}: text='{text}', bbox=[{x1:.1f},{y1:.1f},{x2:.1f},{y2:.1f}]")
                    
                    print(f"DEBUG: Successfully converted {len(formatted_res)} lines!")
                    result = [{
                        'type': 'text',
                        'bbox': [0, 0, w, h],
                        'res': formatted_res
                    }]
                else:
                    print(f"DEBUG: Unexpected first_item type: {type(first_item)}")
                    result = []
            else:
                print(f"DEBUG: result[0] is not a list or is empty")
                result = []
        
        elif isinstance(result, list) and len(result) > 0 and not isinstance(result[0], dict):
             print("DEBUG: Converting list PaddleOCR format to PPStructure format")
             lines = result
             
             formatted_res = []
             for i, line in enumerate(lines):
                 if isinstance(line, (list, tuple)) and len(line) >= 2:
                    box = line[0]
                    text = line[1][0] if isinstance(line[1], (list, tuple)) else str(line[1])
                    pts = np.array(box)
                    x1, y1 = np.min(pts, axis=0)
                    x2, y2 = np.max(pts, axis=0)
                    formatted_res.append({
                        'text': text,
                        'bbox': [float(x1), float(y1), float(x2), float(y2)]
                    })
             
             result = [{
                 'type': 'text',
                 'bbox': [0, 0, w, h],
                 'res': formatted_res
             }]

        # Sort regions by y-coordinate to ensure reading order
        print(f"DEBUG: About to sort. result type = {type(result)}, length = {len(result) if isinstance(result, list) else 'N/A'}")
        if isinstance(result, list) and len(result) > 0:
            print(f"DEBUG: result[0] = {result[0] if isinstance(result[0], dict) else 'NOT DICT'}")
        sorted_res = sorted(result, key=lambda x: (x['bbox'][1], x['bbox'][0]))

        processed_output = []
        for region in sorted_res:
            region_type = region.get('type', '').lower()
            res = region.get('res', {})
            
            item = {
                'type': region_type,
                'bbox': region.get('bbox'), # [x1, y1, x2, y2]
                'html': None,
                'lines': []
            }
            
            # Helper to get bbox from text_region or bbox key
            def get_bbox(data):
                if not isinstance(data, dict):
                    return None
                
                # Try common keys
                box = data.get('bbox') or data.get('text_region') or data.get('poly')
                if not box:
                    return None
                
                # If it's a polygon [[x,y], [x,y], [x,y], [x,y]]
                if isinstance(box, list) and len(box) == 4 and isinstance(box[0], list):
                    pts = np.array(box)
                    x1, y1 = np.min(pts, axis=0)
                    x2, y2 = np.max(pts, axis=0)
                    return [float(x1), float(y1), float(x2), float(y2)]
                
                # If it's a flat list [x1, y1, x2, y2]
                if isinstance(box, list) and len(box) == 4 and not isinstance(box[0], list):
                    return [float(i) for i in box]
                
                return None

            if region_type == 'table':
                item['html'] = res.get('html')
                # Table cells/lines
                table_res = res.get('cell', []) or res.get('content', [])
                if isinstance(table_res, list):
                    for cell in table_res:
                        item['lines'].append({
                            'text': cell.get('text', '') if isinstance(cell, dict) else str(cell),
                            'bbox': get_bbox(cell)
                        })
            else:
                # Regular text region
                if isinstance(res, list):
                    for line in res:
                        if isinstance(line, dict):
                            item['lines'].append({
                                'text': line.get('text', ''),
                                'bbox': get_bbox(line)
                            })
                        elif isinstance(line, (list, tuple)) and len(line) >= 2:
                            # PaddleOCR format: [bbox, (text, conf)]
                            box = line[0]
                            text = line[1][0] if isinstance(line[1], (list, tuple)) else str(line[1])
                            l_item = {'text': text, 'bbox': None}
                            # Convert box to [x1, y1, x2, y2]
                            if isinstance(box, list) and len(box) == 4 and isinstance(box[0], list):
                                pts = np.array(box)
                                x1, y1 = np.min(pts, axis=0)
                                x2, y2 = np.max(pts, axis=0)
                                l_item['bbox'] = [float(x1), float(y1), float(x2), float(y2)]
                            item['lines'].append(l_item)
            
            # Helper to calculate IoU for overlap filtering
            def get_iou(box1, box2):
                if not box1 or not box2: return 0
                x_left = max(box1[0], box2[0])
                y_top = max(box1[1], box2[1])
                x_right = min(box1[2], box2[2])
                y_bottom = min(box1[3], box2[3])
                if x_right < x_left or y_bottom < y_top: return 0.0
                intersection_area = (x_right - x_left) * (y_bottom - y_top)
                area1 = (box1[2] - box1[0]) * (box1[3] - box1[1])
                area2 = (box2[2] - box2[0]) * (box2[3] - box2[1])
                return intersection_area / float(area1 + area2 - intersection_area)

            # Filter out overlapping lines (EXTREMELY aggressive threshold)
            unique_lines = []
            raw_lines = item['lines']
            # Sort by confidence/size if possible, but here we just use order
            for i, line in enumerate(raw_lines):
                if not line.get('bbox'):
                    unique_lines.append(line)
                    continue
                is_duplicate = False
                for existing in unique_lines:
                    if existing.get('bbox') and get_iou(line['bbox'], existing['bbox']) > 0.15:
                        if len(line['text']) > len(existing['text']):
                            existing['text'] = line['text'] # Keep longer text
                        is_duplicate = True
                        break
                if not is_duplicate:
                    unique_lines.append(line)
            
            item['lines'] = unique_lines
            processed_output.append(item)
            
        # --- NEW: Layout Recovery (Absolute Positioning Docx) ---
        docx_path = None
        custom_docx_path = os.path.join(save_folder, f"{img_name}_layout.docx")
        
        # 1. Try Native PaddleOCR Recovery (Better for Tables/Structure)
        try:
            print(f"Attempting native docx recovery for {img_name}...")
            convert_info_docx(img, result, save_folder, img_name)
            
            # Helper: convert_info_docx saves as "{img_name}_ocr.docx"
            built_in_file = os.path.join(save_folder, f"{img_name}_ocr.docx")
            if os.path.exists(built_in_file):
                docx_path = built_in_file
                print(f"✓ Native recovery successful: {docx_path}")
        except Exception as e:
            print(f"Native recovery failed: {e}")

        # 2. Fallback to Custom Layout (if native failed or didn't produce file)
        if not docx_path:
            try:
                print("Falling back to custom layout engine...")
                self.custom_convert_to_docx(processed_output, custom_docx_path, img_width_px=w, img_height_px=h)
                docx_path = custom_docx_path
            except Exception as e:
                print(f"Custom layout docx generation also failed: {e}")

        # Prepare metadata for frontend scaling
        metadata = {
            'width': w,
            'height': h,
            'image_name': img_name
        }
            
        return {
            "processed_output": processed_output,
            "raw_result": result,
            "docx_path": docx_path,
            "metadata": metadata
        }

    def regenerate_docx_from_result(self, result, img_path_or_array, save_folder="./output", img_name="edited_result"):
        """
        Regenerates the Word document using the existing result structure (but potentially edited text).
        Reuses the native convert_info_docx function to ensure 100% consistent formatting.
        """
        if isinstance(img_path_or_array, str):
            img = cv2.imread(img_path_or_array)
        else:
            img = img_path_or_array
            
        if img is None:
            raise ValueError("Image could not be loaded for regeneration.")

        print(f"Regenerating docx for {img_name} with edited data...")
        docx_path = None
        
        # 1. Try Native PaddleOCR Recovery first (Primary method)
        try:
            # result structure MUST retain the original format: 
            # [{'type': 'text', 'bbox': [x,y,w,h], 'res': [{'text': '...', 'bbox': [...]}, ...]}]
            convert_info_docx(img, result, save_folder, img_name)
            
            # Helper: convert_info_docx saves as "{img_name}_ocr.docx"
            built_in_file = os.path.join(save_folder, f"{img_name}_ocr.docx")
            if os.path.exists(built_in_file):
                docx_path = built_in_file
                print(f"✓ Regeneration successful: {docx_path}")
        except Exception as e:
            print(f"Regeneration failed: {e}")
            return None
            
        return docx_path

if __name__ == "__main__":
    # Test initialization
    engine = LocalOCREngine()
    print("Engine Test Complete.")
